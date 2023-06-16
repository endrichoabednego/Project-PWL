# import argparse
# import json
# import re
# from docx import Document
# from docx.enum.shape import WD_INLINE_SHAPE
# from docx.oxml.ns import nsdecls
# from docx.oxml import parse_xml




# # Define command line arguments
# parser = argparse.ArgumentParser()
# parser.add_argument("-d", "--docx", required=True, help="Input DOCX file")
# parser.add_argument("-i", "--indo", required=True, help="Input JSON file for Indonesian translation")
# parser.add_argument("-e", "--eng", required=True, help="Input JSON file for English translation")
# parser.add_argument("-o", "--output", required=True, help="Output DOCX file")
# args = parser.parse_args()

# # Load input files
# with open(args.docx, "rb") as f:
#     document = Document(f)

# with open(args.indo,encoding="utf-8") as f:
#     data_indo = json.load(f)

# with open(args.eng,encoding="utf-8") as f:
#     data_eng = json.load(f)

# # Extract sentences from JSON files
# # regex_span = r"<span id='\d+'>(.+?)</span>"
# regex_span = r"<span id='\d+'>(.+?)</span>"
# kalimat = []
# line=[]
# sentences = []
# baris=[]
# id=[]
# index=[]

# for q in data_indo["sentences"]:
#     q = q.strip()
#     matches = re.split(regex_span, q)
#     if matches:
#         for match in matches:
#             line.append(match+"</span>")
# for q in line:
#     matches = q.split("</span>")
#     for match in matches:
#         kalimat.append(match)



# for q in data_eng["sentences"]:
#     q = q.strip()
#     matches = re.split(regex_span, q)
#     if matches:
#         for match in matches:
#             baris.append(match+"</span>")

# for q in baris:
#     matches = q.split("</span>")
#     for match in matches:
#         sentences.append(match)

# counter1=0
# for q in data_eng["sentences"]:
#     matches1 = re.findall(r"<span id='(\d+)'>.+?</span>", q)
#     if matches1:
#         for match in matches1:
#             id.append(int(match))
#         counter1+=1

# xmlIndex=[]
# indexId=0
# for p in document.paragraphs:
#     inline = p.runs
#     for i in range(len(inline)):
#         xmlIndex.append(i)
#         if i not in id:
#             index.append(i)
#         elif i in id:
#             index.append(id[indexId])
#             indexId+=1

# # # Extract the images from the old document
# # extracted_images = docx2txt.process(args.docx, '')

# for paragraph in document.paragraphs:
#     for run in paragraph.runs:
#         counter = 0
#         for q in kalimat:
#             if run.text == q:
#                 if 'Graphic' in paragraph._p.xml:
#                     continue
#                 elif 'Caption' in paragraph._p.xml:
#                     run.text = ''
#                     new_run = paragraph.add_run(text=sentences[counter], style=run.style)
#                     new_run.bold = run.bold
#                     new_run.italic = run.italic
#                     new_run.underline = run.underline
#                     new_run.font.color.rgb = run.font.color.rgb
#                     new_run.font.size = run.font.size
#                     new_run.font.name = run.font.name
#                 elif 'Hyperlink' in paragraph._p.xml:
#                     run.text = ''
#                     new_run = paragraph.add_run(text=sentences[counter], style=run.style)
#                     new_run.bold = run.bold
#                     new_run.italic = run.italic
#                     new_run.underline = run.underline
#                     new_run.font.color.rgb = run.font.color.rgb
#                     new_run.font.size = run.font.size
#                     new_run.font.name = run.font.name
 
#                 else:
#                     new_run = paragraph.add_run(text=sentences[counter], style=run.style)
#                     new_run.bold = run.bold
#                     new_run.italic = run.italic
#                     new_run.underline = run.underline
#                     new_run.font.color.rgb = run.font.color.rgb
#                     new_run.font.size = run.font.size
#                     new_run.font.name = run.font.name
#                     run.text = ''
#             counter += 1



# for table in document.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 for i, run in enumerate(paragraph.runs):
#                     for j, q in enumerate(kalimat):
#                         if run.text == q:
#                             old_run = paragraph.runs[index[i]]
#                             new_run = paragraph.add_run(text=sentences[j], style=old_run.style)
#                             new_run.bold = old_run.bold
#                             new_run.italic = old_run.italic
#                             new_run.underline = old_run.underline
#                             new_run.font.color.rgb = old_run.font.color.rgb
#                             new_run.font.size = old_run.font.size
#                             new_run.font.name = old_run.font.name
#                             run.clear()
#                             break


# # print('\n\n')
# # print('Index on the xml file: ',xmlIndex)
# # print('ID based on span in english.json file:',id)
# # print('The new index style for the new runs: ',index)
# # print('\n\n\n')
# # print(kalimat)
# # print(sentences)

# # Save output file
# document.save(args.output)
# print(f"File {args.output} has been created.")



import argparse
import json
import re
from docx import Document
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Define command line arguments
parser = argparse.ArgumentParser()
parser.add_argument("-d", "--docx", required=True, help="Input DOCX file")
parser.add_argument("-i", "--indo", required=True, help="Input JSON file for Indonesian translation")
parser.add_argument("-e", "--eng", required=True, help="Input JSON file for English translation")
parser.add_argument("-o", "--output", required=True, help="Output DOCX file")
args = parser.parse_args()

# Load input files
with open(args.docx, "rb") as f:
    document = Document(f)

with open(args.indo, encoding="utf-8") as f:
    data_indo = json.load(f)

with open(args.eng, encoding="utf-8") as f:
    data_eng = json.load(f)

# Extract sentences from JSON files
regex_span = r"<span id='\d+'>(.+?)</span>"
kalimat = [match for q in data_indo["sentences"] for match in re.split(regex_span, q) if match]
sentences = [match for q in data_eng["sentences"] for match in re.split(regex_span, q) if match]
id_matches = [int(match) for q in data_eng["sentences"] for match in re.findall(r"<span id='(\d+)'>.+?</span>", q)]

# Mapping for old index to new index
index_mapping = {i: id_matches[i] if i in id_matches else i for i in range(len(document.paragraphs[0].runs))}


# Helper function to find hyperlink run in a paragraph
def find_hyperlink_run(paragraph):
    for run in paragraph.runs:
        for child in run.element.iter():
            if child.tag.endswith('</w:hyperlink>'):
                return run
    return None


# Replace text in paragraphs and tables
for paragraph in document.paragraphs:
    hyperlink_run = find_hyperlink_run(paragraph)
    if hyperlink_run:
        print('true')
        hyperlink_element = hyperlink_run.element
        hyperlink_id = hyperlink_element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        hyperlink_text = hyperlink_run.text
        if hyperlink_text in kalimat:
            index = kalimat.index(hyperlink_text)
            print(sentences[index])
            new_text = sentences[index]
            new_run = paragraph.add_run(new_text)
            
            new_hyperlink_element = parse_xml(f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" r:id="{hyperlink_id}" w:tgtFrame="_blank" w:history="1"/>')
            new_run._r.insert(0, new_hyperlink_element)
            
            paragraph.runs.remove(hyperlink_run)
    
    for run in paragraph.runs:
        if run.text in kalimat:
            index = kalimat.index(run.text)
            new_text = sentences[index]
            run.text = new_text


# Replace text in tables
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                hyperlink_run = find_hyperlink_run(paragraph)
                if hyperlink_run:
                    hyperlink_element = hyperlink_run.element
                    hyperlink_id = hyperlink_element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    hyperlink_text = hyperlink_run.text

                    if hyperlink_text in kalimat:
                        index = kalimat.index(hyperlink_text)
                        new_text = sentences[index]

                        new_hyperlink_element = parse_xml(f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" r:id="{hyperlink_id}" w:tgtFrame="_blank" w:history="1"/>')
                        new_run = paragraph.add_run(new_text)
                        new_run._r.append(new_hyperlink_element)

                        paragraph.runs.remove(hyperlink_run)
                    else:
                        # Replace the hyperlink text
                        hyperlink_run.text = "New Hyperlink Text"


# Save output file
document.save(args.output)
print(f"File {args.output} has been created.")

